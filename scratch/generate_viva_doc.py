from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_viva_doc(output_path):
    doc = Document()

    # Title
    title = doc.add_heading('HealthMatrix: The Simplified Guide to Heart AI & Math', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_heading('Introduction: What is HealthMatrix?', level=1)
    p = doc.add_paragraph(
        "Imagine you're going on a road trip. A GPS doesn't just show you the map; it looks at traffic, road conditions, "
        "and your speed to tell you when you'll arrive. HealthMatrix is like a GPS for your heart. Instead of traffic, "
        "it looks at your blood pressure, cholesterol, and age to tell you where your health is heading."
    )
    
    # Machine Learning Section
    doc.add_heading('1. The Predictor: How the Computer "Guesses" Your Risk', level=1)
    doc.add_paragraph(
        "In your viva, the teachers will ask about 'Machine Learning.' In plain English, Machine Learning is just a "
        "computer looking at thousands of examples of people with heart disease and learning what they have in common."
    )
    
    doc.add_heading('The Logistic Regression (The Weight-Scale)', level=2)
    doc.add_paragraph(
        "We use something called 'Logistic Regression.' Think of it like a weight-scale (or a seesaw). On one side, we "
        "put all your health data. "
    )
    doc.add_paragraph(
        "- Some things 'weigh' more than others. For example, smoking is a 'heavy' risk factor. It pushes the scale "
        "towards 'High Risk' much faster than being one year older does.\n"
        "- The computer uses a 'Sigmoid Function,' which is just a fancy way of saying a 'Dimmer Switch.' Instead of just "
        "saying YES or NO, it gives a smooth score from 0% to 100%."
    )

    doc.add_heading('The Ensemble (The Team Effort)', level=2)
    doc.add_paragraph(
        "Sometimes computers make mistakes. So, we use an 'Ensemble.' This is like having a Computer and a Human Doctor "
        "talk to each other. The computer looks for patterns in the numbers, but we've also programmed it with 'Human Rules' "
        "(like: 'If blood pressure is over 140, that is always dangerous'). The final score is a blend of both."
    )

    # ECG Section
    doc.add_heading('2. The ECG: Listening to the Heart’s Rhythm', level=1)
    doc.add_paragraph(
        "An ECG (Electrocardiogram) is just a recording of the electrical 'spark' that makes your heart beat. "
        "Every beat has a story:"
    )
    doc.add_paragraph(
        "- The P-wave: The top part of the heart squeezing.\n"
        "- The QRS-complex: The big squeeze (the main pump).\n"
        "- The T-wave: The heart resetting itself for the next beat."
    )
    
    doc.add_heading('How we find "Chaos" (AFib)', level=2)
    doc.add_paragraph(
        "A healthy heart beats like a steady metronome (Tick... Tick... Tick...). But in a condition called Atrial Fibrillation (AFib), "
        "the heart becomes 'chaotic' (Tick..TickTick....Tick). "
    )
    doc.add_paragraph(
        "We use Math to measure this chaos. Specifically, we look at the 'Standard Deviation' of the time between beats. "
        "If the gap between beats is always changing, the computer flags it as 'Irregular' and warns the doctor."
    )

    # Statistics Section
    doc.add_heading('3. Statistics for Humans', level=1)
    doc.add_paragraph(
        "Statistics is the science of understanding 'The Crowd.' We use several tools to understand your health data:"
    )
    
    doc.add_heading('The Shape of Data (Moments)', level=2)
    doc.add_paragraph(
        "1. Mean (The Average): The middle of the pack.\n"
        "2. Variance (The Spread): Do all patients have similar scores, or are they all over the place?\n"
        "3. Skewness (The Lean): If your scores are 'skewed,' it means most people are healthy, but a few are very sick.\n"
        "4. Kurtosis (The Peak): Does your data have extreme outliers? Are there 'freak' events that we need to watch out for?"
    )

    # Bayes Section
    doc.add_heading('4. The Mystery of Bayes Theorem', level=1)
    doc.add_paragraph(
        "Suppose you have a 99% accurate test for a very rare disease. If the test comes back positive, do you "
        "definitely have the disease? Surprisingly, the answer is often 'No!'"
    )
    doc.add_paragraph(
        "This is Bayes' Theorem. It teaches the computer to account for 'Prevalence.' If a disease is super rare, "
        "even a great test will produce a lot of false alarms. We use this math to make sure we don't 'scare' patients "
        "unnecessarily."
    )

    # Regression and Correlation
    doc.add_heading('5. Relationship Math (Correlation & Regression)', level=1)
    
    doc.add_heading('Correlation: Do things happen together?', level=2)
    doc.add_paragraph(
        "If we see that people who exercise more have lower blood pressure, we call that a 'Negative Correlation.' "
        "One goes up, the other goes down. We use the 'Pearson' and 'Spearman' tests to measure how strong this connection is."
    )

    doc.add_heading('Regression: Predicting the Future', level=2)
    doc.add_paragraph(
        "Regression is like 'Connecting the Dots.' If we know how your blood pressure has changed over the last 5 years, "
        "we can draw a straight line into the future to guess where it will be next year. This is 'Forecasting.'"
    )

    # Conclusion
    doc.add_heading('Conclusion: Why this matters', level=1)
    doc.add_paragraph(
        "HealthMatrix isn't just about code; it's about using math to save lives. It takes messy human data "
        "and turns it into clear, simple warnings that a doctor can use to act before it's too late."
    )

    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == '__main__':
    import os
    # Save in the same dir as the script or artifacts
    target = os.path.join(os.getcwd(), 'HealthMatrix_Viva_Preparation.docx')
    create_viva_doc(target)
